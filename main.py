import os
import docx
import pandas as pd
import xlsxwriter


def main():
    directory = os.getcwd()
    files = os.listdir(directory)
    docs = list(filter(lambda x: x.endswith('.docx'), files))
    # print(images)
    table_for_block = [[], [], [], [], [], [], []]
    for file in docs:
        doc = docx.Document(file)
        tables = []
        i = 0
        Text2 = []
        for table in doc.tables:
            tables.append(table)
            for index, row in enumerate(table.rows):
                if i == 1:
                    i += 1
                    break
                if index == 0:
                    row_text = list(cell.text for cell in row.cells)
                for cell in row.cells:
                    # if "(" not in cell.text:
                    # if cell.text != " " and cell.text != "":
                    if i == 0 and cell.text == "A":
                        i += 1
                        break
                    Text2.append(cell.text)
        all_tables = doc.tables
        data_tables = {i: None for i in range(len(all_tables))}
        for i, table in enumerate(all_tables):
            # создаем список строк для таблицы `i` (пока пустые)
            # print(i)
            data_tables[i] = [[] for _ in range(len(table.rows))]
            # проходимся по строкам таблицы `i`
            for j, row in enumerate(table.rows):
                # проходимся по ячейкам таблицы `i` и строки `j`
                # print("")
                for cell in row.cells:
                    # добавляем значение ячейки в соответствующий
                    # список, созданного словаря под данные таблиц
                    data_tables[i][j].append(cell.text)
                # if i==6:

                    # print(data_tables[i][j])
                    # print("-"*20)

        for j in range(len(data_tables[6])):
            table_for_block[0].append(Text2[5])
            for k in range(len(data_tables[6][j])):
                table_for_block[k+1].append(data_tables[6][j][k])
        # print(table_for_block[0])
        # print(len(table_for_block[0]))
    name_block = ["Код","Описание трудовых функций, входящих в профессиональный стандарт (функциональная карта вида профессиональной деятельности)",""," ","  ","   ","    "]
    dict_for_block = dict(zip(name_block, table_for_block))
    # print(dict_for_block)
    data_frameBlock = pd.DataFrame(dict_for_block)

    # print(data_frameBlock)
    excelwriter = pd.ExcelWriter('Профстандарты.xlsx', engine='xlsxwriter')
    data_frameBlock.to_excel(excelwriter, sheet_name='Лист1', index=False)
    excelwriter.save()

if __name__ == '__main__':
    main()